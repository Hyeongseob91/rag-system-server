"""File Parsers"""
import json
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

import pdfplumber
from docx import Document
from openpyxl import load_workbook

from ...schemas.preprocessing import RawDocument


class BaseParser(ABC):
    @property
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        pass

    @abstractmethod
    def parse(self, file_path: str) -> RawDocument:
        pass

    def can_you_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower().lstrip(".")
        return ext in self.supported_extensions

    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        stat = path.stat()
        return {
            "file_name": path.name,
            "file_type": path.suffix.lower().lstrip("."),
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }


class PDFParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]:
        return ["pdf"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
                for table in page.extract_tables():
                    if table:
                        table_text = "\n".join(" | ".join(str(c or "") for c in row) for row in table)
                        if table_text not in text:
                            pages[-1] += f"\n\n[Table]\n{table_text}"
            file_info["page_count"] = len(pdf.pages)
        return RawDocument(
            content="\n\n".join(pages), source=str(Path(file_path).absolute()),
            file_type="pdf", file_name=file_info["file_name"], metadata=file_info, pages=pages
        )


class DOCXParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]:
        return ["docx"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            paragraphs.append(f"[Table]\n" + "\n".join(rows))
        file_info["paragraph_count"] = len(doc.paragraphs)
        return RawDocument(
            content="\n\n".join(paragraphs), source=str(Path(file_path).absolute()),
            file_type="docx", file_name=file_info["file_name"], metadata=file_info, pages=paragraphs
        )


class XLSXParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]:
        return ["xlsx", "xls"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        wb = load_workbook(file_path, read_only=True, data_only=True)
        sheets, all_content = {}, []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = [" | ".join(str(c) if c else "" for c in row) for row in ws.iter_rows(values_only=True) if any(row)]
            content = "\n".join(rows)
            if content.strip():
                sheets[sheet_name] = content
                all_content.append(f"[Sheet: {sheet_name}]\n{content}")
        wb.close()
        file_info["sheet_count"] = len(wb.sheetnames)
        return RawDocument(
            content="\n\n".join(all_content), source=str(Path(file_path).absolute()),
            file_type="xlsx", file_name=file_info["file_name"], metadata=file_info, sheets=sheets
        )


class TXTParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]:
        return ["txt", "md", "rst"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        content = None
        for enc in ["utf-8", "cp949", "euc-kr", "latin-1"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    content = f.read()
                file_info["encoding"] = enc
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            raise ValueError(f"파일 인코딩 감지 실패: {file_path}")
        return RawDocument(
            content=content, source=str(Path(file_path).absolute()),
            file_type=file_info["file_type"], file_name=file_info["file_name"], metadata=file_info
        )


class JSONParser(BaseParser):
    @property
    def supported_extensions(self) -> List[str]:
        return ["json"]

    def parse(self, file_path: str) -> RawDocument:
        file_info = self._get_file_info(file_path)
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        content = self._json_to_text(data)
        return RawDocument(
            content=content, source=str(Path(file_path).absolute()),
            file_type="json", file_name=file_info["file_name"], metadata=file_info
        )

    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        lines = []
        if isinstance(data, dict):
            for k, v in data.items():
                if isinstance(v, (dict, list)):
                    lines.append(f"{prefix}{k}:")
                    lines.append(self._json_to_text(v, prefix + "  "))
                else:
                    lines.append(f"{prefix}{k}: {v}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, prefix + "  "))
                else:
                    lines.append(f"{prefix}- {item}")
        else:
            lines.append(f"{prefix}{data}")
        return "\n".join(lines)


class UnifiedFileParser:
    def __init__(self):
        self._parsers = [PDFParser(), DOCXParser(), XLSXParser(), TXTParser(), JSONParser()]

    def get_supported_extensions(self) -> List[str]:
        return [ext for p in self._parsers for ext in p.supported_extensions]

    def parse(self, file_path: str) -> RawDocument:
        for parser in self._parsers:
            if parser.can_you_parse(file_path):
                return parser.parse(file_path)
        raise ValueError(f"지원하지 않는 파일 형식: {Path(file_path).suffix}")
